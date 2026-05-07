from docx import Document

doc = Document('format.docx')

text = ''

for para in doc.paragraphs:

    text += para.text + '\n'

for table in doc.tables:

    for row in table.rows:

        for cell in row.cells:

            text += cell.text + '\n'

print(text)